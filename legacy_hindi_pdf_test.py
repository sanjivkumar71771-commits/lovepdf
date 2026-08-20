#!/usr/bin/env python3
"""
Test script for LEGACY (non-Unicode) Hindi font PDF -> Word conversion.

Tests the fix for PDFs authored with legacy Devanagari fonts (Kruti Dev / DevLys / Chanakya)
that store Hindi as ASCII-mapped glyph codes, causing gibberish extraction.

The fix detects legacy font names and force-OCRs the rendered pages to recover
proper Unicode Devanagari.
"""
import os
import sys
import requests
import zipfile
from pathlib import Path
from fpdf import FPDF
from docx import Document

# Backend URL from environment
BACKEND_URL = os.getenv('REACT_APP_BACKEND_URL', 'https://lovepdf-tools-1.preview.emergentagent.com')
API_BASE = f"{BACKEND_URL}/api"

# Font path
LOHIT_FONT = "/usr/share/fonts/truetype/lohit-devanagari/Lohit-Devanagari.ttf"

# Test data
HINDI_TEXT_LINES = [
    "चरित्र प्रमाण पत्र",
    "प्रमाणित किया जाता है कि श्री कुमारी श्रीमती",
    "इनका नैतिक चरित्र उत्तम है",
    "यह प्रमाण पत्र उनके अनुरोध पर जारी किया गया है",
]

ENGLISH_TEXT_LINES = [
    "Character Certificate",
    "This is to certify that Mr./Ms./Mrs.",
    "Their moral character is excellent",
    "This certificate is issued upon their request",
]


def create_legacy_hindi_pdf(output_path: Path) -> Path:
    """
    Create a PDF using fpdf with a LEGACY FONT NAME (KrutiDev010) but actual
    Lohit-Devanagari.ttf font, writing real Devanagari text.
    
    This simulates a legacy Hindi PDF: the font name contains 'kruti' (triggering
    legacy detection), but the visual rendering is correct Devanagari.
    """
    pdf = FPDF()
    pdf.add_page()
    
    # Register the Lohit font with a LEGACY NAME containing 'kruti'
    pdf.add_font('KrutiDev010', '', LOHIT_FONT, uni=True)
    pdf.set_font('KrutiDev010', size=24)
    
    # Write real Devanagari text
    for line in HINDI_TEXT_LINES:
        pdf.cell(0, 15, txt=line, ln=True)
    
    pdf.output(str(output_path))
    return output_path


def create_normal_hindi_pdf(output_path: Path) -> Path:
    """
    Create a PDF with a NON-LEGACY font name (NotoDeva) using the same Lohit font,
    writing Devanagari text. This should use the normal text-layer path, not OCR.
    """
    pdf = FPDF()
    pdf.add_page()
    
    # Register with a NON-LEGACY name
    pdf.add_font('NotoDeva', '', LOHIT_FONT, uni=True)
    pdf.set_font('NotoDeva', size=24)
    
    # Write real Devanagari text
    for line in HINDI_TEXT_LINES:
        pdf.cell(0, 15, txt=line, ln=True)
    
    pdf.output(str(output_path))
    return output_path


def create_english_pdf(output_path: Path) -> Path:
    """Create a simple English text PDF for regression testing."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=16)
    
    for line in ENGLISH_TEXT_LINES:
        pdf.cell(0, 10, txt=line, ln=True)
    
    pdf.output(str(output_path))
    return output_path


def is_valid_docx(path: Path) -> bool:
    """Check if a file is a valid .docx (valid zip with required parts)."""
    try:
        if not path.exists() or path.stat().st_size < 200:
            return False
        if not zipfile.is_zipfile(str(path)):
            return False
        with zipfile.ZipFile(str(path)) as z:
            names = set(z.namelist())
            required = ["[Content_Types].xml", "word/document.xml"]
            if not set(required).issubset(names):
                return False
            if z.testzip() is not None:
                return False
        # Try to open with python-docx
        doc = Document(str(path))
        _ = doc.paragraphs
        return True
    except Exception as e:
        print(f"  ❌ DOCX validation failed: {e}")
        return False


def extract_docx_text(path: Path) -> str:
    """Extract all text from a .docx file."""
    try:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        print(f"  ❌ Text extraction failed: {e}")
        return ""


def contains_devanagari(text: str) -> bool:
    """Check if text contains Devanagari Unicode characters (U+0900–U+097F)."""
    return any('\u0900' <= c <= '\u097f' for c in text)


def count_non_whitespace(text: str) -> int:
    """Count non-whitespace characters."""
    return len(''.join(text.split()))


def test_pdf_to_word(pdf_path: Path, test_name: str, lang: str = "eng", 
                     expect_devanagari: bool = False) -> dict:
    """
    POST a PDF to /api/pdf/pdf-to-word and validate the response.
    
    Returns a dict with test results.
    """
    print(f"\n{'='*70}")
    print(f"TEST: {test_name}")
    print(f"{'='*70}")
    
    result = {
        "test_name": test_name,
        "success": False,
        "http_status": None,
        "docx_size": 0,
        "text_length": 0,
        "has_devanagari": False,
        "text_preview": "",
        "error": None,
    }
    
    try:
        # POST the PDF
        url = f"{API_BASE}/pdf/pdf-to-word"
        with open(pdf_path, 'rb') as f:
            files = {'file': (pdf_path.name, f, 'application/pdf')}
            data = {'lang': lang}
            print(f"  📤 POST {url}")
            print(f"     File: {pdf_path.name} ({pdf_path.stat().st_size} bytes)")
            print(f"     Lang: {lang}")
            
            response = requests.post(url, files=files, data=data, timeout=120)
            result["http_status"] = response.status_code
            print(f"  📥 Response: {response.status_code}")
        
        if response.status_code != 200:
            result["error"] = f"HTTP {response.status_code}: {response.text[:200]}"
            print(f"  ❌ FAILED: {result['error']}")
            return result
        
        # Save the response as a .docx
        docx_path = pdf_path.parent / f"{pdf_path.stem}_output.docx"
        docx_path.write_bytes(response.content)
        result["docx_size"] = docx_path.stat().st_size
        print(f"  💾 Saved output: {docx_path.name} ({result['docx_size']} bytes)")
        
        # Validate the .docx
        if not is_valid_docx(docx_path):
            result["error"] = "Invalid .docx file (not a valid zip or missing required parts)"
            print(f"  ❌ FAILED: {result['error']}")
            return result
        print(f"  ✅ Valid .docx (zip structure + python-docx compatible)")
        
        # Extract text
        text = extract_docx_text(docx_path)
        result["text_length"] = count_non_whitespace(text)
        result["has_devanagari"] = contains_devanagari(text)
        result["text_preview"] = text[:200] if text else "(empty)"
        
        print(f"  📝 Extracted text length: {result['text_length']} non-whitespace chars")
        print(f"  🔤 Contains Devanagari: {result['has_devanagari']}")
        print(f"  📄 Text preview: {result['text_preview'][:100]}...")
        
        # Check if empty
        if result["text_length"] == 0:
            result["error"] = "DOCX is empty (no text extracted)"
            print(f"  ❌ FAILED: {result['error']}")
            return result
        
        # Check Devanagari expectation
        if expect_devanagari and not result["has_devanagari"]:
            result["error"] = "Expected Devanagari Unicode chars but found none (likely gibberish)"
            print(f"  ❌ FAILED: {result['error']}")
            return result
        
        result["success"] = True
        print(f"  ✅ PASSED")
        return result
        
    except Exception as e:
        result["error"] = str(e)
        print(f"  ❌ EXCEPTION: {e}")
        return result


def main():
    print("\n" + "="*70)
    print("LEGACY HINDI FONT PDF -> WORD CONVERSION TEST SUITE")
    print("="*70)
    
    # Create temp directory for test files
    test_dir = Path("/tmp/legacy_hindi_test")
    test_dir.mkdir(exist_ok=True)
    print(f"\n📁 Test directory: {test_dir}")
    
    results = []
    
    # TEST 1: Legacy Hindi font (MAIN TEST)
    print("\n" + "🔴"*35)
    print("TEST 1: LEGACY HINDI FONT (Kruti Dev simulation)")
    print("🔴"*35)
    print("Creating PDF with font name 'KrutiDev010' (contains 'kruti' token)")
    print("This should trigger legacy font detection and force OCR path")
    
    legacy_pdf = create_legacy_hindi_pdf(test_dir / "legacy_hindi.pdf")
    print(f"✅ Created: {legacy_pdf.name} ({legacy_pdf.stat().st_size} bytes)")
    
    result1 = test_pdf_to_word(
        legacy_pdf, 
        "Legacy Hindi Font (KrutiDev010) -> OCR Recovery",
        lang="hin",
        expect_devanagari=True
    )
    results.append(result1)
    
    # TEST 2: English regression
    print("\n" + "🟢"*35)
    print("TEST 2: ENGLISH REGRESSION")
    print("🟢"*35)
    
    english_pdf = create_english_pdf(test_dir / "english.pdf")
    print(f"✅ Created: {english_pdf.name} ({english_pdf.stat().st_size} bytes)")
    
    result2 = test_pdf_to_word(
        english_pdf,
        "English Text PDF -> Word",
        lang="eng",
        expect_devanagari=False
    )
    results.append(result2)
    
    # TEST 3: Normal Unicode Hindi (non-legacy font name)
    print("\n" + "🟡"*35)
    print("TEST 3: NORMAL UNICODE HINDI (Non-legacy font)")
    print("🟡"*35)
    print("Creating PDF with font name 'NotoDeva' (no legacy token)")
    print("This should use normal text-layer path, not OCR")
    
    normal_hindi_pdf = create_normal_hindi_pdf(test_dir / "normal_hindi.pdf")
    print(f"✅ Created: {normal_hindi_pdf.name} ({normal_hindi_pdf.stat().st_size} bytes)")
    
    result3 = test_pdf_to_word(
        normal_hindi_pdf,
        "Normal Unicode Hindi (NotoDeva) -> Word",
        lang="hin",
        expect_devanagari=True
    )
    results.append(result3)
    
    # Summary
    print("\n" + "="*70)
    print("TEST SUMMARY")
    print("="*70)
    
    passed = sum(1 for r in results if r["success"])
    total = len(results)
    
    for i, r in enumerate(results, 1):
        status = "✅ PASS" if r["success"] else "❌ FAIL"
        print(f"\n{i}. {r['test_name']}")
        print(f"   Status: {status}")
        print(f"   HTTP: {r['http_status']}")
        print(f"   DOCX size: {r['docx_size']} bytes")
        print(f"   Text length: {r['text_length']} chars")
        print(f"   Has Devanagari: {r['has_devanagari']}")
        if r["error"]:
            print(f"   Error: {r['error']}")
        if r["text_preview"]:
            print(f"   Preview: {r['text_preview'][:80]}...")
    
    print("\n" + "="*70)
    print(f"OVERALL: {passed}/{total} tests passed")
    print("="*70)
    
    # Core pass criterion
    if result1["success"] and result1["has_devanagari"]:
        print("\n🎉 CORE CRITERION MET: Legacy Hindi font PDF yields real Devanagari (not gibberish)")
    else:
        print("\n❌ CORE CRITERION FAILED: Legacy Hindi font PDF did not produce valid Devanagari")
    
    return 0 if passed == total else 1


if __name__ == "__main__":
    sys.exit(main())
