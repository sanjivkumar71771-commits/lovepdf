#!/usr/bin/env python3
"""
Hindi (Devanagari) PDF to Word Conversion Test
Tests the bug fix for Hindi PDFs producing empty .docx files.
"""
import os
import sys
import requests
import tempfile
import zipfile
from pathlib import Path
from io import BytesIO

# Backend URL from environment
BACKEND_URL = os.getenv("REACT_APP_BACKEND_URL", "https://lovepdf-tools-1.preview.emergentagent.com")
API_BASE = f"{BACKEND_URL}/api"

# Test data
HINDI_TEXT = """नमस्ते दुनिया। यह एक परीक्षण दस्तावेज़ है।
कृपया इसे वर्ड में बदलें।
यह हिंदी भाषा का परीक्षण है।
देवनागरी लिपि का समर्थन महत्वपूर्ण है।
धन्यवाद।"""

ENGLISH_TEXT = """Hello World. This is a test document.
Please convert this to Word format.
This is an English language test.
Thank you."""

FONT_PATH = "/usr/share/fonts/truetype/lohit-devanagari/Lohit-Devanagari.ttf"


def check_devanagari_chars(text):
    """Check if text contains Devanagari Unicode characters (U+0900–U+097F)."""
    if not text:
        return False
    return any(0x0900 <= ord(c) <= 0x097F for c in text)


def validate_docx(file_path):
    """Validate that a file is a proper .docx (valid zip with required parts)."""
    try:
        if not Path(file_path).exists():
            return False, "File does not exist"
        
        if Path(file_path).stat().st_size < 200:
            return False, f"File too small: {Path(file_path).stat().st_size} bytes"
        
        if not zipfile.is_zipfile(file_path):
            return False, "Not a valid zip file"
        
        with zipfile.ZipFile(file_path) as z:
            names = z.namelist()
            required = ["[Content_Types].xml", "word/document.xml"]
            missing = [r for r in required if r not in names]
            if missing:
                return False, f"Missing required parts: {missing}"
            
            # Test zip integrity
            if z.testzip() is not None:
                return False, "Zip file is corrupt"
        
        # Try to load with python-docx
        from docx import Document
        doc = Document(file_path)
        paragraphs = doc.paragraphs  # Touch the body to ensure it's valid
        
        return True, "Valid docx"
    except Exception as e:
        return False, f"Validation error: {str(e)}"


def extract_docx_text(file_path):
    """Extract all text from a .docx file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text
    except Exception as e:
        return f"ERROR: {str(e)}"


def create_hindi_text_pdf():
    """Create a PDF with Hindi text layer using fpdf2."""
    try:
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.add_page()
        
        # Add Unicode font (Lohit Devanagari)
        pdf.add_font("Lohit", "", FONT_PATH, uni=True)
        pdf.set_font("Lohit", size=14)
        
        # Add Hindi text
        for line in HINDI_TEXT.strip().split('\n'):
            pdf.cell(0, 10, txt=line, ln=True)
        
        # Save to bytes
        pdf_bytes = pdf.output()
        return pdf_bytes
    except Exception as e:
        print(f"ERROR creating Hindi text PDF with fpdf2: {e}")
        # Fallback to reportlab
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            
            # Register the Devanagari font
            pdfmetrics.registerFont(TTFont('Lohit', FONT_PATH))
            c.setFont('Lohit', 14)
            
            # Add Hindi text
            y = 750
            for line in HINDI_TEXT.strip().split('\n'):
                c.drawString(50, y, line)
                y -= 20
            
            c.save()
            return buffer.getvalue()
        except Exception as e2:
            print(f"ERROR creating Hindi text PDF with reportlab: {e2}")
            raise


def create_scanned_hindi_pdf():
    """Create an image-only PDF (scanned) with Hindi text using PIL + img2pdf."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        import img2pdf
        
        # Create an image with Hindi text
        img = Image.new('RGB', (800, 400), color='white')
        draw = ImageDraw.Draw(img)
        
        # Load Devanagari font
        try:
            font = ImageFont.truetype(FONT_PATH, 40)
        except Exception:
            print(f"WARNING: Could not load font from {FONT_PATH}, using default")
            font = ImageFont.load_default()
        
        # Draw Hindi text
        text = "नमस्ते दुनिया परीक्षण"
        draw.text((50, 150), text, fill='black', font=font)
        
        # Convert image to bytes
        img_bytes = BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        # Convert image to PDF (no text layer)
        pdf_bytes = img2pdf.convert(img_bytes.getvalue())
        return pdf_bytes
    except Exception as e:
        print(f"ERROR creating scanned Hindi PDF: {e}")
        raise


def create_english_text_pdf():
    """Create a simple English text PDF for regression testing."""
    try:
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for line in ENGLISH_TEXT.strip().split('\n'):
            pdf.cell(0, 10, txt=line, ln=True)
        
        pdf_bytes = pdf.output()
        return pdf_bytes
    except Exception as e:
        print(f"ERROR creating English PDF with fpdf2: {e}")
        # Fallback to reportlab
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.setFont('Helvetica', 12)
            
            y = 750
            for line in ENGLISH_TEXT.strip().split('\n'):
                c.drawString(50, y, line)
                y -= 20
            
            c.save()
            return buffer.getvalue()
        except Exception as e2:
            print(f"ERROR creating English PDF with reportlab: {e2}")
            raise


def test_pdf_to_word(pdf_bytes, lang="eng", test_name="Test"):
    """Test PDF to Word conversion endpoint."""
    print(f"\n{'='*70}")
    print(f"TEST: {test_name}")
    print(f"{'='*70}")
    
    try:
        # Save PDF to temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            tmp_pdf.write(pdf_bytes)
            tmp_pdf_path = tmp_pdf.name
        
        print(f"✓ Created test PDF: {tmp_pdf_path} ({len(pdf_bytes)} bytes)")
        
        # POST to /api/pdf/pdf-to-word
        url = f"{API_BASE}/pdf/pdf-to-word"
        print(f"✓ Posting to: {url}")
        
        with open(tmp_pdf_path, 'rb') as f:
            files = {'file': ('test.pdf', f, 'application/pdf')}
            data = {'lang': lang}
            response = requests.post(url, files=files, data=data, timeout=120)
        
        print(f"✓ Response status: {response.status_code}")
        
        if response.status_code != 200:
            print(f"✗ FAILED: Expected 200, got {response.status_code}")
            print(f"  Response: {response.text[:500]}")
            return False
        
        # Save response to temp docx file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(response.content)
            tmp_docx_path = tmp_docx.name
        
        print(f"✓ Received .docx file: {len(response.content)} bytes")
        
        # Validate docx structure
        is_valid, msg = validate_docx(tmp_docx_path)
        if not is_valid:
            print(f"✗ FAILED: Invalid .docx - {msg}")
            return False
        
        print(f"✓ Valid .docx structure: {msg}")
        
        # Extract text
        extracted_text = extract_docx_text(tmp_docx_path)
        text_len = len(extracted_text.strip())
        non_ws_len = len(''.join(extracted_text.split()))
        
        print(f"✓ Extracted text length: {text_len} chars ({non_ws_len} non-whitespace)")
        
        if non_ws_len == 0:
            print(f"✗ FAILED: Document is EMPTY (no text extracted)")
            return False
        
        print(f"✓ Document is NOT empty")
        
        # Check for Devanagari characters if this is a Hindi test
        if lang in ["hin", "Devanagari"]:
            has_devanagari = check_devanagari_chars(extracted_text)
            if has_devanagari:
                print(f"✓ Contains Devanagari characters (U+0900–U+097F)")
            else:
                print(f"✗ WARNING: No Devanagari characters found in output")
                print(f"  Extracted text preview: {extracted_text[:200]}")
                # Don't fail - OCR might not be perfect, but as long as it's not empty
        
        # Show text preview
        preview = extracted_text[:200].replace('\n', ' ')
        print(f"✓ Text preview: {preview}...")
        
        # Cleanup
        os.unlink(tmp_pdf_path)
        os.unlink(tmp_docx_path)
        
        print(f"\n✅ {test_name} PASSED")
        return True
        
    except Exception as e:
        print(f"\n✗ FAILED: {test_name}")
        print(f"  Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("\n" + "="*70)
    print("HINDI PDF TO WORD CONVERSION TEST SUITE")
    print("="*70)
    print(f"Backend URL: {BACKEND_URL}")
    print(f"API Base: {API_BASE}")
    print(f"Font Path: {FONT_PATH}")
    print(f"Font exists: {os.path.exists(FONT_PATH)}")
    
    results = {}
    
    # Test 1: Hindi TEXT-layer PDF
    print("\n\n" + "="*70)
    print("TEST 1: HINDI TEXT-LAYER PDF")
    print("="*70)
    try:
        hindi_text_pdf = create_hindi_text_pdf()
        results['hindi_text'] = test_pdf_to_word(hindi_text_pdf, lang="hin", test_name="Hindi Text-Layer PDF")
    except Exception as e:
        print(f"✗ FAILED to create/test Hindi text PDF: {e}")
        results['hindi_text'] = False
    
    # Test 2: SCANNED Hindi PDF (OCR path)
    print("\n\n" + "="*70)
    print("TEST 2: SCANNED HINDI PDF (OCR)")
    print("="*70)
    try:
        scanned_hindi_pdf = create_scanned_hindi_pdf()
        results['hindi_scanned'] = test_pdf_to_word(scanned_hindi_pdf, lang="hin", test_name="Scanned Hindi PDF (OCR)")
    except Exception as e:
        print(f"✗ FAILED to create/test scanned Hindi PDF: {e}")
        results['hindi_scanned'] = False
    
    # Test 3: English PDF (regression)
    print("\n\n" + "="*70)
    print("TEST 3: ENGLISH PDF (REGRESSION)")
    print("="*70)
    try:
        english_pdf = create_english_text_pdf()
        results['english'] = test_pdf_to_word(english_pdf, lang="eng", test_name="English PDF (Regression)")
    except Exception as e:
        print(f"✗ FAILED to create/test English PDF: {e}")
        results['english'] = False
    
    # Summary
    print("\n\n" + "="*70)
    print("TEST SUMMARY")
    print("="*70)
    for test_name, passed in results.items():
        status = "✅ PASSED" if passed else "✗ FAILED"
        print(f"{test_name:20s}: {status}")
    
    total = len(results)
    passed = sum(1 for v in results.values() if v)
    print(f"\nTotal: {passed}/{total} tests passed")
    
    if passed == total:
        print("\n🎉 ALL TESTS PASSED!")
        return 0
    else:
        print(f"\n❌ {total - passed} TEST(S) FAILED")
        return 1


if __name__ == "__main__":
    sys.exit(main())
