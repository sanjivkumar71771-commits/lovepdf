#!/usr/bin/env python3
"""
PDF Conversion Tools Backend API Tests
Tests all PDF conversion endpoints in /api/pdf
"""
import requests
import sys
import os
import tempfile
from pathlib import Path
import zipfile

# Backend URL from frontend/.env
BASE_URL = "https://lovepdf-tools-1.preview.emergentagent.com/api/pdf"

# Test results tracking
passed = 0
failed = 0
test_results = []

def log_test(name, success, details=""):
    global passed, failed
    if success:
        passed += 1
        status = "✅ PASS"
    else:
        failed += 1
        status = "❌ FAIL"
    msg = f"{status}: {name}"
    if details:
        msg += f" - {details}"
    print(msg)
    test_results.append({"name": name, "success": success, "details": details})


def create_text_pdf():
    """Create a simple text-based PDF using fpdf2"""
    from fpdf import FPDF
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Sample Text PDF Document", ln=True)
    pdf.cell(200, 10, txt="This is a test document with readable text.", ln=True)
    pdf.cell(200, 10, txt="It should be converted to Word format successfully.", ln=True)
    pdf.cell(200, 10, txt="The conversion should preserve the text content.", ln=True)
    pdf.output(temp_file.name)
    
    return temp_file.name


def create_table_pdf():
    """Create a PDF with a table using fpdf2"""
    from fpdf import FPDF
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=14, style='B')
    pdf.cell(200, 10, txt="Sales Report - Q1 2024", ln=True)
    pdf.ln(5)
    
    # Create table header
    pdf.set_font("Arial", size=12, style='B')
    pdf.cell(60, 10, txt="Product", border=1)
    pdf.cell(40, 10, txt="Quantity", border=1)
    pdf.cell(40, 10, txt="Price", border=1)
    pdf.ln()
    
    # Table data
    pdf.set_font("Arial", size=12)
    data = [
        ['Widget A', '10', '$25.00'],
        ['Widget B', '5', '$50.00'],
        ['Widget C', '15', '$15.00'],
        ['Total', '30', '$90.00']
    ]
    
    for row in data:
        pdf.cell(60, 10, txt=row[0], border=1)
        pdf.cell(40, 10, txt=row[1], border=1)
        pdf.cell(40, 10, txt=row[2], border=1)
        pdf.ln()
    
    pdf.output(temp_file.name)
    return temp_file.name


def create_scanned_pdf():
    """Create an image-only PDF (simulating a scanned document) with NO text layer"""
    from PIL import Image, ImageDraw, ImageFont
    import img2pdf
    
    # Create an image with text rendered as pixels
    img = Image.new('RGB', (800, 1000), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except Exception:
        font = ImageFont.load_default()
    
    # Draw text as image (not as PDF text layer)
    draw.text((50, 50), "Scanned Document Test", fill='black', font=font)
    draw.text((50, 100), "This document has no text layer.", fill='black', font=font)
    draw.text((50, 150), "OCR should recognize this text.", fill='black', font=font)
    draw.text((50, 200), "Important keywords: invoice payment receipt", fill='black', font=font)
    
    # Save as image first
    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    img.save(temp_img.name, "PNG")
    
    # Convert image to PDF (no text layer)
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    with open(temp_pdf.name, "wb") as f:
        f.write(img2pdf.convert(temp_img.name))
    
    os.unlink(temp_img.name)
    return temp_pdf.name


def create_docx():
    """Create a simple .docx file"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test Word document.')
    doc.add_paragraph('It should be converted to PDF format.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Some content in section 1.')
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def validate_docx(file_path):
    """Strictly validate a .docx file"""
    path = Path(file_path)
    
    # Check file exists and has size
    if not path.exists() or path.stat().st_size < 200:
        return False, "File missing or too small"
    
    # Check it's a valid zip
    if not zipfile.is_zipfile(str(path)):
        return False, "Not a valid zip file"
    
    # Check required OOXML parts
    try:
        with zipfile.ZipFile(str(path)) as z:
            names = set(z.namelist())
            required = ["[Content_Types].xml", "word/document.xml"]
            if not set(required).issubset(names):
                return False, f"Missing required parts: {set(required) - names}"
            
            # Check for corruption
            if z.testzip() is not None:
                return False, "Zip file is corrupt"
    except Exception as e:
        return False, f"Zip validation error: {e}"
    
    # Try to open with python-docx
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = doc.paragraphs
        
        # Check it has content
        if len(paragraphs) == 0:
            return False, "Document has no paragraphs"
        
        # Check there's actual text
        text = "\n".join([p.text for p in paragraphs])
        if len(text.strip()) == 0:
            return False, "Document has no text content"
        
        return True, f"Valid docx with {len(paragraphs)} paragraphs, {len(text)} chars"
    except Exception as e:
        return False, f"python-docx validation error: {e}"


def validate_xlsx(file_path):
    """Strictly validate an .xlsx file"""
    path = Path(file_path)
    
    # Check file exists and has size
    if not path.exists() or path.stat().st_size < 200:
        return False, "File missing or too small"
    
    # Check it's a valid zip
    if not zipfile.is_zipfile(str(path)):
        return False, "Not a valid zip file"
    
    # Check required OOXML parts
    try:
        with zipfile.ZipFile(str(path)) as z:
            names = set(z.namelist())
            required = ["[Content_Types].xml", "xl/workbook.xml"]
            if not set(required).issubset(names):
                return False, f"Missing required parts: {set(required) - names}"
            
            # Check for corruption
            if z.testzip() is not None:
                return False, "Zip file is corrupt"
    except Exception as e:
        return False, f"Zip validation error: {e}"
    
    # Try to open with openpyxl
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True)
        sheets = wb.sheetnames
        
        if len(sheets) == 0:
            wb.close()
            return False, "Workbook has no sheets"
        
        # Check at least one sheet has data
        has_data = False
        for sheet_name in sheets:
            ws = wb[sheet_name]
            for row in ws.iter_rows(max_row=10):
                for cell in row:
                    if cell.value is not None and str(cell.value).strip():
                        has_data = True
                        break
                if has_data:
                    break
        
        wb.close()
        
        if not has_data:
            return False, "No data found in any sheet"
        
        return True, f"Valid xlsx with {len(sheets)} sheet(s)"
    except Exception as e:
        return False, f"openpyxl validation error: {e}"


def validate_pdf(file_path):
    """Validate a PDF file"""
    path = Path(file_path)
    
    if not path.exists() or path.stat().st_size < 100:
        return False, "File missing or too small"
    
    # Check PDF header
    with open(path, 'rb') as f:
        header = f.read(5)
        if not header.startswith(b'%PDF-'):
            return False, "Invalid PDF header"
    
    return True, f"Valid PDF ({path.stat().st_size} bytes)"


def test_health():
    """Test 1: GET /api/pdf/health"""
    print("\n" + "="*60)
    print("TEST 1: PDF TOOLS HEALTH CHECK")
    print("="*60)
    
    try:
        resp = requests.get(f"{BASE_URL}/health", timeout=10)
        
        if resp.status_code == 200:
            data = resp.json()
            if "tools" in data:
                tools = data["tools"]
                required_tools = ["soffice", "gs", "qpdf", "tesseract", "pdftoppm", "ocrmypdf"]
                
                all_true = all(tools.get(t, False) for t in required_tools)
                
                if all_true:
                    log_test("GET /api/pdf/health", True, f"All tools available: {', '.join(required_tools)}")
                else:
                    missing = [t for t in required_tools if not tools.get(t, False)]
                    log_test("GET /api/pdf/health", False, f"Missing tools: {', '.join(missing)}")
            else:
                log_test("GET /api/pdf/health", False, "Missing 'tools' in response")
        else:
            log_test("GET /api/pdf/health", False, f"Status {resp.status_code}: {resp.text}")
    except Exception as e:
        log_test("GET /api/pdf/health", False, str(e))


def test_pdf_to_word_text():
    """Test 2: POST /api/pdf/pdf-to-word with TEXT-based PDF"""
    print("\n" + "="*60)
    print("TEST 2: PDF TO WORD (TEXT-BASED PDF)")
    print("="*60)
    
    pdf_path = None
    output_path = None
    
    try:
        # Create text PDF
        pdf_path = create_text_pdf()
        print(f"Created text PDF: {pdf_path}")
        
        # Upload and convert
        with open(pdf_path, 'rb') as f:
            files = {'file': ('test.pdf', f, 'application/pdf')}
            resp = requests.post(f"{BASE_URL}/pdf-to-word", files=files, timeout=60)
        
        if resp.status_code == 200:
            # Save response
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            with open(output_path, 'wb') as f:
                f.write(resp.content)
            
            print(f"Received .docx: {output_path} ({len(resp.content)} bytes)")
            
            # Strict validation
            is_valid, msg = validate_docx(output_path)
            
            if is_valid:
                log_test("POST /api/pdf/pdf-to-word (text PDF)", True, msg)
            else:
                log_test("POST /api/pdf/pdf-to-word (text PDF)", False, f"Invalid docx: {msg}")
        else:
            log_test("POST /api/pdf/pdf-to-word (text PDF)", False, f"Status {resp.status_code}: {resp.text[:300]}")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-word (text PDF)", False, str(e))
    finally:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if output_path and os.path.exists(output_path):
            os.unlink(output_path)


def test_pdf_to_excel():
    """Test 3: POST /api/pdf/pdf-to-excel with TABLE PDF"""
    print("\n" + "="*60)
    print("TEST 3: PDF TO EXCEL (TABLE PDF)")
    print("="*60)
    
    pdf_path = None
    output_path = None
    
    try:
        # Create table PDF
        pdf_path = create_table_pdf()
        print(f"Created table PDF: {pdf_path}")
        
        # Upload and convert
        with open(pdf_path, 'rb') as f:
            files = {'file': ('table.pdf', f, 'application/pdf')}
            resp = requests.post(f"{BASE_URL}/pdf-to-excel", files=files, timeout=60)
        
        if resp.status_code == 200:
            # Save response
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
            with open(output_path, 'wb') as f:
                f.write(resp.content)
            
            print(f"Received .xlsx: {output_path} ({len(resp.content)} bytes)")
            
            # Strict validation
            is_valid, msg = validate_xlsx(output_path)
            
            if is_valid:
                log_test("POST /api/pdf/pdf-to-excel (table PDF)", True, msg)
            else:
                log_test("POST /api/pdf/pdf-to-excel (table PDF)", False, f"Invalid xlsx: {msg}")
        else:
            log_test("POST /api/pdf/pdf-to-excel (table PDF)", False, f"Status {resp.status_code}: {resp.text[:300]}")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-excel (table PDF)", False, str(e))
    finally:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if output_path and os.path.exists(output_path):
            os.unlink(output_path)


def test_pdf_to_word_scanned():
    """Test 4: POST /api/pdf/pdf-to-word with SCANNED/image-only PDF (OCR path)"""
    print("\n" + "="*60)
    print("TEST 4: PDF TO WORD (SCANNED/IMAGE-ONLY PDF - OCR)")
    print("="*60)
    
    pdf_path = None
    output_path = None
    
    try:
        # Create scanned PDF (image-only, no text layer)
        pdf_path = create_scanned_pdf()
        print(f"Created scanned PDF: {pdf_path}")
        
        # Upload and convert
        with open(pdf_path, 'rb') as f:
            files = {'file': ('scanned.pdf', f, 'application/pdf')}
            resp = requests.post(f"{BASE_URL}/pdf-to-word", files=files, timeout=120)
        
        if resp.status_code == 200:
            # Save response
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            with open(output_path, 'wb') as f:
                f.write(resp.content)
            
            print(f"Received .docx: {output_path} ({len(resp.content)} bytes)")
            
            # Strict validation
            is_valid, msg = validate_docx(output_path)
            
            if not is_valid:
                log_test("POST /api/pdf/pdf-to-word (scanned PDF)", False, f"Invalid docx: {msg}")
            else:
                # Check OCR worked - should contain some of the text we rendered
                from docx import Document
                doc = Document(output_path)
                text = "\n".join([p.text for p in doc.paragraphs]).lower()
                
                # Look for keywords we rendered in the image
                keywords = ["scanned", "document", "ocr", "text"]
                found_keywords = [kw for kw in keywords if kw in text]
                
                if len(found_keywords) >= 2:
                    log_test("POST /api/pdf/pdf-to-word (scanned PDF)", True, 
                            f"Valid docx with OCR text. Found keywords: {', '.join(found_keywords)}")
                else:
                    log_test("POST /api/pdf/pdf-to-word (scanned PDF)", False, 
                            f"OCR may have failed. Only found: {', '.join(found_keywords) if found_keywords else 'none'}")
        else:
            log_test("POST /api/pdf/pdf-to-word (scanned PDF)", False, f"Status {resp.status_code}: {resp.text[:300]}")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-word (scanned PDF)", False, str(e))
    finally:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if output_path and os.path.exists(output_path):
            os.unlink(output_path)


def test_office_to_pdf():
    """Test 5: POST /api/pdf/office-to-pdf with .docx file"""
    print("\n" + "="*60)
    print("TEST 5: OFFICE TO PDF (.docx to PDF)")
    print("="*60)
    
    docx_path = None
    output_path = None
    
    try:
        # Create docx
        docx_path = create_docx()
        print(f"Created .docx: {docx_path}")
        
        # Upload and convert
        with open(docx_path, 'rb') as f:
            files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            resp = requests.post(f"{BASE_URL}/office-to-pdf", files=files, timeout=60)
        
        if resp.status_code == 200:
            # Save response
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
            with open(output_path, 'wb') as f:
                f.write(resp.content)
            
            print(f"Received PDF: {output_path} ({len(resp.content)} bytes)")
            
            # Validate PDF
            is_valid, msg = validate_pdf(output_path)
            
            if is_valid:
                log_test("POST /api/pdf/office-to-pdf (.docx)", True, msg)
            else:
                log_test("POST /api/pdf/office-to-pdf (.docx)", False, f"Invalid PDF: {msg}")
        else:
            log_test("POST /api/pdf/office-to-pdf (.docx)", False, f"Status {resp.status_code}: {resp.text[:300]}")
    except Exception as e:
        log_test("POST /api/pdf/office-to-pdf (.docx)", False, str(e))
    finally:
        if docx_path and os.path.exists(docx_path):
            os.unlink(docx_path)
        if output_path and os.path.exists(output_path):
            os.unlink(output_path)


def main():
    print("="*60)
    print("PDF CONVERSION TOOLS BACKEND API TESTS")
    print("="*60)
    print(f"Backend URL: {BASE_URL}")
    print("="*60)
    
    # Run all tests
    test_health()
    test_pdf_to_word_text()
    test_pdf_to_excel()
    test_pdf_to_word_scanned()
    test_office_to_pdf()
    
    # Summary
    print("\n" + "="*60)
    print("TEST SUMMARY")
    print("="*60)
    print(f"Total: {passed + failed} tests")
    print(f"✅ Passed: {passed}")
    print(f"❌ Failed: {failed}")
    print("="*60)
    
    if failed > 0:
        print("\nFAILED TESTS:")
        for result in test_results:
            if not result["success"]:
                print(f"  - {result['name']}: {result['details']}")
        sys.exit(1)
    else:
        print("\n🎉 All PDF conversion tests passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
